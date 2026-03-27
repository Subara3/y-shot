"""Convert y-shot manual markdown to Word (.docx) with embedded images."""
import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(DOCS_DIR, "y-shot_manual.md")
OUT_PATH = os.path.join(DOCS_DIR, "y-shot_manual.docx")

def parse_md():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        return f.read()

def add_image(doc, img_name):
    img_path = os.path.join(DOCS_DIR, img_name)
    if os.path.isfile(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic UI"
    style.font.size = Pt(10)

    md = parse_md()
    lines = md.split("\n")

    in_table = False
    table_headers = []
    table_rows = []
    in_code = False
    code_lines = []

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if table_headers or table_rows:
            add_table(doc, table_headers, table_rows)
            doc.add_paragraph()
        in_table = False
        table_headers = []
        table_rows = []

    def flush_code():
        nonlocal in_code, code_lines
        if code_lines:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.left_indent = Inches(0.3)
        in_code = False
        code_lines = []

    for line in lines:
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code:
                flush_code()
            else:
                if in_table: flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        # Table
        if "|" in stripped and stripped.startswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if all(set(c) <= set("- :") for c in cells):
                continue  # separator row
            if not in_table:
                in_table = True
                table_headers = [re.sub(r"\*\*(.+?)\*\*", r"\1", c) for c in cells]
            else:
                table_rows.append([re.sub(r"\*\*(.+?)\*\*", r"\1", c) for c in cells])
            continue
        else:
            if in_table: flush_table()

        # Skip horizontal rules
        if stripped == "---":
            continue

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title = stripped[2:].strip()
            p = doc.add_heading(title, level=0)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue

        # Images
        img_match = re.match(r"!\[.*?\]\((.+?)\)", stripped)
        if img_match:
            add_image(doc, img_match.group(1))
            continue

        # Blockquotes
        if stripped.startswith("> "):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Empty lines
        if not stripped:
            continue

        # Normal text (with bold handling)
        text = stripped
        # Clean up markdown formatting
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            bold_match = re.match(r"\*\*(.+?)\*\*", part)
            if bold_match:
                run = p.add_run(bold_match.group(1))
                run.bold = True
            else:
                p.add_run(part)

    # Flush remaining
    if in_table: flush_table()
    if in_code: flush_code()

    doc.save(OUT_PATH)
    print(f"Generated: {OUT_PATH}")

if __name__ == "__main__":
    build()
